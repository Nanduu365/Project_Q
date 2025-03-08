import sys
import cv2
import time
import wave
import pyaudio
import os
import numpy as np
import mediapipe as mp
from tensorflow.keras.models import load_model
from gtts import gTTS
import pygame
from faster_whisper import WhisperModel
from langchain_groq import ChatGroq
import shutil
from docx import Document
from docx.shared import Inches
import tempfile

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QComboBox,
    QPushButton, QLabel, QTextEdit, QGroupBox, QFileDialog, QMessageBox,
    QStackedWidget,QSplashScreen, QScrollArea, QGridLayout
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, pyqtSlot
from PyQt6.QtGui import QImage, QPixmap, QFont

# Global Variables

model = None
character_dict = None
mp_hands = None
hands_for_sign = None

audio_to_text_model = None

main_llm = None

tts = None
#all these global variables will be updated when the main window is loading

# Camera Thread for sign Preview
class CameraThread(QThread):
    changePixmap = pyqtSignal(QImage)
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self._run_flag = True
        self.last_frame = None  # Store the most recent RGB frame

    def run(self):
        cap = cv2.VideoCapture(0)
        while self._run_flag:
            ret, frame = cap.read()
            if ret:
                rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                self.last_frame = rgb_frame.copy()
                h, w, ch = rgb_frame.shape
                bytes_per_line = ch * w
                qt_img = QImage(rgb_frame.data, w, h, bytes_per_line, QImage.Format.Format_RGB888)
                self.changePixmap.emit(qt_img)
            time.sleep(0.03)
        cap.release()

    def stop(self):
        self._run_flag = False
        self.wait()

#Audio Recorder
class AudioRecorder(QThread):
    def __init__(self, filename="input_audio.wav", parent=None):
        super().__init__(parent)
        self.filename = filename
        self._recording = True

    def run(self):
        CHUNK = 1024
        FORMAT = pyaudio.paInt16
        CHANNELS = 1
        RATE = 44100

        p = pyaudio.PyAudio()
        stream = p.open(format=FORMAT,
                        channels=CHANNELS,
                        rate=RATE,
                        input=True,
                        frames_per_buffer=CHUNK)
        frames = []
        while self._recording:
            try:
                data = stream.read(CHUNK)
                frames.append(data)
            except Exception as e:
                print("Audio recording error:", e)
                break
        stream.stop_stream()
        stream.close()
        p.terminate()

        wf = wave.open(self.filename, 'wb')
        wf.setnchannels(CHANNELS)
        wf.setsampwidth(p.get_sample_size(FORMAT))
        wf.setframerate(RATE)
        wf.writeframes(b''.join(frames))
        wf.close()

    def stop(self):
        self._recording = False


class AudioPlayerThread(QThread):
    progressChanged = pyqtSignal(int)
    finishedPlaying = pyqtSignal()
    
    def __init__(self, filename):
        super().__init__()
        self.filename = filename

    def run(self):
        
        if self.filename.split('.')[1] == 'wav': 
        
            try:
                wf = wave.open(self.filename, 'rb')
            except FileNotFoundError:
                self.finishedPlaying.emit()
                return

            total_frames = wf.getnframes()
            frames_played = 0

            p = pyaudio.PyAudio()
            stream = p.open(format=p.get_format_from_width(wf.getsampwidth()),
                            channels=wf.getnchannels(),
                            rate=wf.getframerate(),
                            output=True)
            CHUNK = 1024
            data = wf.readframes(CHUNK)
            while data:
                stream.write(data)
                frames_played += CHUNK
                progress = int((frames_played / total_frames) * 100)
                self.progressChanged.emit(min(progress, 100))
                data = wf.readframes(CHUNK)
            stream.stop_stream()
            stream.close()
            p.terminate()
            wf.close()
            self.progressChanged.emit(100)

        else:
            pygame.init()
            pygame.mixer.init()
            pygame.mixer.music.load(self.filename)
            pygame.mixer.music.play()

        self.finishedPlaying.emit()

#sign Player Thread (for LLM sign output)
class SignPlayerThread(QThread):
    frameChanged = pyqtSignal(QImage)
    finishedPlaying = pyqtSignal()
    
    def __init__(self, text, folder="Alphabets", parent=None):
        super().__init__(parent)
        self.text = text
        self.folder = folder
        self._run_flag = True

    def run(self):
        # Process each character in the response text.
        for letter in self.text:
            if not self._run_flag:
                break
            if letter == ' ':
                path = os.path.join(self.folder, "space.png")
            else:
                path = os.path.join(self.folder, letter.upper() + ".png")
            # Load image as QImage (thread-safe)
            image = QImage(path)
            if not image.isNull():
                self.frameChanged.emit(image)
            self.msleep(250)  # 250 ms per frame
        self.finishedPlaying.emit()

    def stop(self):
        self._run_flag = False
        self.wait()


class AudioInputWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.is_recording = False
        self.recorder = None
        self.audio_player = None
        self.transcribed_text = ""
        group = QGroupBox("Audio Input")
        layout = QVBoxLayout()
        self.record_button = QPushButton("Start Recording")
        self.record_button.clicked.connect(self.toggle_recording)
        layout.addWidget(self.record_button)
        self.play_recording_button = QPushButton("Play Recording")
        self.play_recording_button.setEnabled(False)
        self.play_recording_button.clicked.connect(self.play_input_audio)
        layout.addWidget(self.play_recording_button)
        self.confirm_recording_button = QPushButton("Confirm Recording")
        self.confirm_recording_button.setEnabled(False)
        self.confirm_recording_button.clicked.connect(self.confirm_recording)
        layout.addWidget(self.confirm_recording_button)
        group.setLayout(layout)
        main_layout = QVBoxLayout(self)
        main_layout.addWidget(group)

    def toggle_recording(self):
        if not self.is_recording:
            self.record_button.setText("Stop Recording")
            self.is_recording = True
            self.play_recording_button.setEnabled(False)
            self.confirm_recording_button.setEnabled(False)
            self.recorder = AudioRecorder("temp_input_audio.wav")
            self.recorder.start()
        else:
            self.record_button.setText("Start Recording")
            self.is_recording = False
            if self.recorder is not None:
                self.recorder.stop()
                self.recorder.wait()
                self.play_recording_button.setEnabled(True)
                self.confirm_recording_button.setEnabled(True)

    def play_input_audio(self):
        filename = "temp_input_audio.wav"
        if self.audio_player is not None and self.audio_player.isRunning():
            return
        self.audio_player = AudioPlayerThread(filename)
        self.audio_player.start()

    def confirm_recording(self):
        ret = QMessageBox.question(
            self,
            "Confirm Recording",
            "Do you want to proceed with this recording?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if ret == QMessageBox.StandardButton.Yes:
            try:
                if os.path.exists("input_audio.wav"):
                    os.remove("input_audio.wav")
                os.rename("temp_input_audio.wav", "input_audio.wav")
                segments,_= audio_to_text_model.transcribe("input_audio.wav")
                text = [segment.text for segment in segments]
                text = ''.join(text)
                self.transcribed_text = text
                self.play_recording_button.setEnabled(False)
                self.confirm_recording_button.setEnabled(False)
            except Exception as e:
                QMessageBox.warning(self, "Error", f"Error saving file: {e}")
        else:
            if os.path.exists("temp_input_audio.wav"):
                os.remove("temp_input_audio.wav")
            QMessageBox.information(self, "Recording", "Recording discarded. Please record again.")
            self.play_recording_button.setEnabled(False)
            self.confirm_recording_button.setEnabled(False)

class SignInputWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.camera_thread = None
        group = QGroupBox("sign Input")
        layout = QVBoxLayout()
        self.camera_label = QLabel("Camera Preview")
        self.camera_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.camera_label)
        group.setLayout(layout)
        main_layout = QVBoxLayout(self)
        main_layout.addWidget(group)

    @pyqtSlot(QImage)
    def set_camera_image(self, qt_img):
        pix = QPixmap.fromImage(qt_img).scaled(self.camera_label.size(), Qt.AspectRatioMode.KeepAspectRatio)
        self.camera_label.setPixmap(pix)

    def showEvent(self, event):
        super().showEvent(event)
        if self.camera_thread is None:
            self.camera_thread = CameraThread()
            self.camera_thread.changePixmap.connect(self.set_camera_image)
            self.camera_thread.start()

    def hideEvent(self, event):
        super().hideEvent(event)
        if self.camera_thread is not None:
            self.camera_thread.stop()
            self.camera_thread = None



class SignOutputWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.sign_thread = None
        self.frames = []  # To store QPixmap frames.
        self.scroll_area = QScrollArea(self)
        self.scroll_area.setWidgetResizable(True)
        self.container = QWidget()
        self.grid_layout = QGridLayout(self.container)
        # Remove spacing and margins from the grid layout.
        self.grid_layout.setSpacing(0)
        self.grid_layout.setContentsMargins(0, 0, 0, 0)
        self.container.setLayout(self.grid_layout)
        self.scroll_area.setWidget(self.container)
        layout = QVBoxLayout(self)
        # Remove spacing and margins from the main layout.
        layout.setSpacing(0)
        layout.setContentsMargins(0, 0, 0, 0)
        title_label = QLabel("sign output")
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title_label)
        layout.addWidget(self.scroll_area)
        self.setLayout(layout)
        self.columns = 10
        self.current_row = 0
        self.current_col = 0

    def play_sign(self, text):
        # Clear previous grid.
        while self.grid_layout.count():
            item = self.grid_layout.takeAt(0)
            widget = item.widget()
            if widget is not None:
                widget.deleteLater()
        self.frames = []
        self.current_row = 0
        self.current_col = 0
        if self.sign_thread is not None:
            self.sign_thread.stop()
        self.sign_thread = SignPlayerThread(text)
        self.sign_thread.frameChanged.connect(self.update_frame)
        self.sign_thread.start()

    @pyqtSlot(QImage)
    def update_frame(self, image):
        # Convert QImage to QPixmap in the main (GUI) thread.
        pixmap = QPixmap.fromImage(image)
        self.frames.append(pixmap)
        label = QLabel()
        label.setPixmap(pixmap)
        # Set a fixed size for each image.
        label.setFixedSize(60, 100)
        # Remove any extra margins or padding.
        label.setContentsMargins(0, 0, 0, 0)
        label.setStyleSheet("margin: 0px; padding: 0px;")
        self.grid_layout.addWidget(label, self.current_row, self.current_col)
        self.current_col += 1
        if self.current_col >= self.columns:
            self.current_col = 0
            self.current_row += 1

    def stop(self):
        if self.sign_thread is not None:
            self.sign_thread.stop()


class SignPredictionThread(QThread):
    newText = pyqtSignal(str)
    def __init__(self, sign_input, parent=None):
        super().__init__(parent)
        self.sign_input = sign_input
        self._run_flag = True
        self.current_text = ""
    def run(self):
        while self._run_flag:
            time.sleep(2)  # sample one frame every 2 seconds
            if self.sign_input.camera_thread is None:
                continue
            frame = self.sign_input.camera_thread.last_frame
            if frame is None:
                continue
            results = hands_for_sign.process(frame)
            landmark_list = []
            if results.multi_hand_landmarks:
                for hand_landmarks in results.multi_hand_landmarks:
                    for landmark in hand_landmarks.landmark:
                        landmark_list.extend([landmark.x, landmark.y, landmark.z])
            landmark_array = np.array(landmark_list)
            if landmark_array.shape[0] == 63:
                prediction = model.predict(landmark_array.reshape(1, -1))
                predicted_class = np.argmax(prediction)
                predicted_char = character_dict[predicted_class]
                if predicted_char == 'space':
                    self.current_text += " "
                elif predicted_char == 'del':
                    self.current_text = self.current_text[:-1]
                elif predicted_char == 'nothing':
                    pass
                else:
                    self.current_text += predicted_char
                self.newText.emit(self.current_text)
    def stop(self):
        self._run_flag = False
        self.wait()


# Text Input and Output Widgets 
class TextInputWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        group = QGroupBox("Text Input")
        layout = QVBoxLayout()
        self.input_text = QTextEdit()
        self.input_text.setPlaceholderText("Enter your text here...")
        layout.addWidget(self.input_text)
        group.setLayout(layout)
        main_layout = QVBoxLayout(self)
        main_layout.addWidget(group)

class TextOutputWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        group = QGroupBox("Text Output")
        layout = QVBoxLayout()
        self.output_text = QTextEdit()
        self.output_text.setReadOnly(True)
        layout.addWidget(self.output_text)
        group.setLayout(layout)
        main_layout = QVBoxLayout(self)
        main_layout.addWidget(group)

#sign Input with Generated Text Widget --
class SignInputWithGeneratedTextWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.current_text = ""
        layout = QHBoxLayout(self)
        self.sign_input = SignInputWidget()
        self.generated_text_group = QGroupBox("Generated Text")
        gen_layout = QVBoxLayout()
        self.generated_text = QTextEdit()
        self.generated_text.setReadOnly(True)
        self.generated_text.setFont(QFont("Courier", 14))
        self.generated_text.setText("")
        gen_layout.addWidget(self.generated_text)
        self.generated_text_group.setLayout(gen_layout)
        layout.addWidget(self.sign_input, 1)
        layout.addWidget(self.generated_text_group, 1)
        
        # Start a dedicated thread for sign prediction.
        self.prediction_thread = SignPredictionThread(self.sign_input)
        self.prediction_thread.newText.connect(self.update_generated_text)
        self.prediction_thread.start()

    @pyqtSlot(str)
    def update_generated_text(self, text):
        self.generated_text.setText(text)

    def hideEvent(self, event):
        super().hideEvent(event)
        if self.prediction_thread is not None:
            self.prediction_thread.stop()

# Combines an input widget and an output widget horizontally 
class MixedModeWidget(QWidget):
    def __init__(self, input_mode, output_mode, parent=None):
        super().__init__(parent)
        self.input_mode = input_mode
        self.output_mode = output_mode
        layout = QHBoxLayout(self)
        self.input_widget = get_input_widget(input_mode)
        self.output_widget = get_output_widget(output_mode)
        if input_mode == "Text" and output_mode != "Text":
            input_stretch = 1
            output_stretch = 2
        elif input_mode == 'Sign' and output_mode == "Text":
            input_stretch = 2
            output_stretch = 1
        elif input_mode == "Audio" and output_mode == "Text":
            input_stretch = 1
            output_stretch = 2
        else:
            input_stretch = 1
            output_stretch = 1
        layout.addWidget(self.input_widget, input_stretch)
        layout.addWidget(self.output_widget, output_stretch)


def get_input_widget(mode):
    if mode == "Text":
        return TextInputWidget()
    elif mode == "Audio":
        return AudioInputWidget()
    elif mode == 'Sign':
        return SignInputWithGeneratedTextWidget()
    else:
        from PyQt6.QtWidgets import QLabel
        return QLabel("Unsupported Input Mode")

def get_output_widget(mode):
    if mode == "Text":
        return TextOutputWidget()
    elif mode == "Audio":
        return AudioOutputWidget()
    elif mode == 'Sign':
        return SignOutputWidget()
    else:
        from PyQt6.QtWidgets import QLabel
        return QLabel("Unsupported Output Mode")

class AudioOutputWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.audio_player = None
        group = QGroupBox("Audio Output")
        layout = QVBoxLayout()
        self.play_output_button = QPushButton("Play Output Audio")
        self.stop_playing_button = QPushButton('Stop Playing')
        self.stop_playing_button.clicked.connect(self.stop_playing)
        self.play_output_button.clicked.connect(self.play_output_audio)
        layout.addWidget(self.play_output_button)
        layout.addWidget(self.stop_playing_button)
        group.setLayout(layout)
        main_layout = QVBoxLayout(self)
        main_layout.addWidget(group)

    def play_output_audio(self):
        filename = "output_audio.mp3"
        if self.audio_player is not None and self.audio_player.isRunning():
            return
        self.audio_player = AudioPlayerThread(filename)
        self.audio_player.start()
    
    def stop_playing(self):
        pygame.mixer.music.stop()
        pygame.mixer.quit()






#Main Application Window
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Muti Chat")
        self.resize(900, 600)
        self.response = None  # Will store the main LLM response
        self.uploaded_file_text = ""  # Store text from uploaded file (if any)
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        # Top layout with dropdowns.
        top_layout = QHBoxLayout()
        self.input_mode_combo = QComboBox()
        self.input_mode_combo.addItems(["Text", "Audio", 'Sign'])
        self.output_mode_combo = QComboBox()
        self.output_mode_combo.addItems(["Text", "Audio", 'Sign'])
        top_layout.addWidget(self.input_mode_combo)
        top_layout.addStretch(1)
        top_layout.addWidget(self.output_mode_combo)
        main_layout.addLayout(top_layout)
        # Dynamic middle area.
        self.stacked_widget = QStackedWidget()
        main_layout.addWidget(self.stacked_widget, stretch=1)
        
        self.input_mode_combo.currentTextChanged.connect(self.update_mode)
        self.output_mode_combo.currentTextChanged.connect(self.update_mode)
        # Bottom layout with buttons.
        bottom_layout = QHBoxLayout()
        self.send_button = QPushButton("Send")
        self.send_button.clicked.connect(self.send_to_llm)
        self.file_uploader_button = QPushButton("Upload File")
        self.file_uploader_button.clicked.connect(self.upload_file)
        self.remove_file_button = QPushButton('Remove File')
        self.remove_file_button.clicked.connect(self.remove_file)
        self.download_button = QPushButton("Download")
        self.download_button.clicked.connect(self.download_file)
        bottom_layout.addWidget(self.send_button)
        bottom_layout.addWidget(self.file_uploader_button)
        bottom_layout.addWidget(self.remove_file_button)
        bottom_layout.addWidget(self.download_button)
        main_layout.addLayout(bottom_layout)
        self.update_mode()
        self.apply_styling()
        


    def update_mode(self):
        input_mode = self.input_mode_combo.currentText()
        output_mode = self.output_mode_combo.currentText()
        widget = MixedModeWidget(input_mode, output_mode)
        
        self.stacked_widget.addWidget(widget)
        self.stacked_widget.setCurrentWidget(widget)



    def send_to_llm(self):
        # Retrieve user text from the current input.
        input_mode = self.input_mode_combo.currentText()
        current_page = self.stacked_widget.currentWidget()
        user_text = ""
        if input_mode == "Text":
            if hasattr(current_page, "input_text"):
                user_text = current_page.input_text.toPlainText()
            elif hasattr(current_page, "input_widget"):
                user_text = current_page.input_widget.input_text.toPlainText()
        elif input_mode == "Audio":

            if hasattr(current_page, "transcribed_text"):
                user_text = current_page.transcribed_text
            elif hasattr(current_page, "input_widget") and hasattr(current_page.input_widget, "transcribed_text"):
                user_text = current_page.input_widget.transcribed_text
        elif input_mode == 'Sign':
            if hasattr(current_page, "generated_text"):
                user_text = current_page.generated_text.toPlainText()
            elif hasattr(current_page, "input_widget") and hasattr(current_page.input_widget, "generated_text"):
                user_text = current_page.input_widget.generated_text.toPlainText()
        # Append uploaded file text if exists.
        if self.uploaded_file_text:
            user_text = user_text + "\n" + self.uploaded_file_text
        if not user_text.strip():
            QMessageBox.warning(self, "Input Error", "No input text available to send.")
            return
        try:
            response = main_llm.invoke(user_text)
            self.response = response
            out_mode = self.output_mode_combo.currentText()
            if out_mode == 'Sign':
                output_widget = None
                if hasattr(current_page, "output_widget"):
                    output_widget = current_page.output_widget
                elif hasattr(current_page, "play_sign"):
                    output_widget = current_page
                if output_widget and hasattr(output_widget, "play_sign"):
                    output_widget.play_sign(response.content)
                else:
                    QMessageBox.information(self, "LLM Response", f"Response:\n{response.content}")
            elif out_mode == "Audio":
            
                tts = gTTS(text = response.content, lang='en')
                tts.save('output_audio.mp3')

                if hasattr(current_page, "output_widget") and hasattr(current_page.output_widget, "play_output_button"):
                    current_page.output_widget.play_output_button.setStyleSheet("background-color: green")
                
            else:
                if hasattr(current_page, "output_text"):
                    current_page.output_text.setPlainText(response.content)
                elif hasattr(current_page, "output_widget") and hasattr(current_page.output_widget, "output_text"):
                    current_page.output_widget.output_text.setPlainText(response.content)
            
        except Exception as e:
            QMessageBox.warning(self, "LLM Error", f"Error invoking LLM: {e}")

    def upload_file(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Upload File", "", "Text Files (*.txt);;Word Documents (*.docx)")
        if file_path.endswith('.txt'):
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    self.uploaded_file_text = f.read()
                QMessageBox.information(self, "File Uploaded", "File uploaded and content loaded.")
            except Exception as e:
                QMessageBox.warning(self, "File Error", f"Error reading file: {e}")
        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            self.uploaded_file_text = "\n".join([para.text for para in doc.paragraphs])
        else:
            self.uploaded_file_text = ""

    def remove_file(self):
        self.uploaded_file_text = ""

    def download_file(self):
        # Download the output shown in the output column based on current output mode.
        out_mode = self.output_mode_combo.currentText()
        current_page = self.stacked_widget.currentWidget()
        if out_mode == "Text":
            # Get text from output text widget.
            content = ""
            if hasattr(current_page, "output_text"):
                content = current_page.output_text.toPlainText()
            elif hasattr(current_page, "output_widget") and hasattr(current_page.output_widget, "output_text"):
                content = current_page.output_widget.output_text.toPlainText()
            if not content.strip():
                QMessageBox.warning(self, "Download Error", "No text output available.")
                return
            save_path, _ = QFileDialog.getSaveFileName(self, "Download Text Output", "output.txt", "Text Files (*.txt)")
            if save_path:
                try:
                    with open(save_path, "w", encoding="utf-8") as f:
                        f.write(content)
                    QMessageBox.information(self, "Download Complete", f"Text output saved to: {save_path}")
                except Exception as e:
                    QMessageBox.warning(self, "Download Error", f"Error saving file: {e}")
        elif out_mode == "Audio":
            # For audio, download the output_audio.wav file.
            if not os.path.exists("output_audio.mp3"):
                QMessageBox.warning(self, "Download Error", "No audio output available.")
                return
            save_path, _ = QFileDialog.getSaveFileName(self, "Download Audio Output", "output_audio.wav", "WAV Files (*.wav)")
            if save_path:
                try:
                    
                    shutil.copyfile("output_audio.mp3", save_path)
                    QMessageBox.information(self, "Download Complete", f"Audio output saved to: {save_path}")
                except Exception as e:
                    QMessageBox.warning(self, "Download Error", f"Error copying file: {e}")
        elif out_mode == 'Sign':
            # For sign, retrieve frames from SignOutputWidget.
            frames = []
            if hasattr(current_page, "output_widget") and isinstance(current_page.output_widget, SignOutputWidget):
                frames = current_page.output_widget.frames
            elif hasattr(current_page, "frames"):
                frames = current_page.frames
            if not frames:
                QMessageBox.warning(self, "Download Error", "No sign frames available for download.")
                return
            save_path, _ = QFileDialog.getSaveFileName(self, "Download sign Output as Word Document", "sign_output.docx", "Word Documents (*.docx)")
            if save_path:
                try:
                    
                    doc = Document()
                    paragraph = doc.add_paragraph()
                    for pixmap in frames:
                        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                        temp_file.close()
                        pixmap.save(temp_file.name, "PNG")

                        run = paragraph.add_run() 
                        run.add_picture(temp_file.name, width=Inches(1.0))
                        os.unlink(temp_file.name)
                    doc.save(save_path)
                    QMessageBox.information(self, "Download Complete", f"sign output saved as Word document: {save_path}")
                except Exception as e:
                    QMessageBox.warning(self, "Download Error", f"Error saving Word document: {e}")
        else:
            QMessageBox.warning(self, "Download Error", "Unsupported output mode.")

    def apply_styling(self):
        with open("theme.qss", "r") as f:
            self.setStyleSheet(f.read())

        



def main():
    app = QApplication(sys.argv)
    splash = QSplashScreen(QPixmap(), Qt.WindowType.WindowStaysOnTopHint)
    splash.setFixedSize(300,120)
    splash.setFont(QFont("Arial",10))
    
    splash.showMessage("Loading models, please wait...", Qt.AlignmentFlag.AlignCenter, Qt.GlobalColor.white)
    splash.show()
    app.processEvents()
    
    global model, character_dict, mp_hands, hands_for_sign, audio_to_text_model, main_llm,ttts
    # Load sign predictor model and related globals.
    model = load_model('sign_predictor.keras')
    character_dict = {
        0: 'A', 1: 'B', 2: 'C', 3: 'D', 4: 'E', 5: 'F', 6: 'G', 7: 'H', 8: 'I', 9: 'J',
        10: 'K', 11: 'L', 12: 'M', 13: 'N', 14: 'O', 15: 'P', 16: 'Q', 17: 'R', 18: 'S', 19: 'T',
        20: 'U', 21: 'V', 22: 'W', 23: 'X', 24: 'Y', 25: 'Z', 26: 'space', 27: 'del', 28: 'nothing'
    }
    mp_hands = mp.solutions.hands
    hands_for_sign = mp_hands.Hands(static_image_mode=True, max_num_hands=1, min_detection_confidence=0.5)
    # Load Whisper audio-to-text model.
    audio_to_text_model = WhisperModel("base")
    # Load main LLM via ChatGroq.
   

    API_KEY = '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'
    main_llm = ChatGroq(
        model='llama-3.1-8b-instant',
        groq_api_key=API_KEY,
        temperature=0,
        max_tokens=250
    )
    
    
    time.sleep(1)
    window = MainWindow()
    splash.finish(window)
    window.show()
    sys.exit(app.exec())

if __name__ == "__main__":
    main()



